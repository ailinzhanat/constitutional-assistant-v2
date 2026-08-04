"""
Constitutional Assistant - File Parser
Поддерживаемые форматы:
- PDF (текст + сканы с OCR)
- Word (.docx, .doc)
- Изображения (JPG, PNG, TIFF, BMP) с OCR
- Текст (.txt)
- OpenDocument (.odt)
"""

import os
import io
import tempfile
from pathlib import Path
from typing import Optional
from fastapi import UploadFile

# ============================================================================
# ОГРАНИЧЕНИЯ БЕЗОПАСНОСТИ
# ============================================================================

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 МБ

# Магические байты для проверки реального типа файла
# (расширение можно подделать, содержимое — нельзя)
MAGIC_BYTES = {
    b"%PDF": "pdf",
    b"PK\x03\x04": "zip",       # docx / odt (оба — ZIP-архивы)
    b"\xff\xd8\xff": "jpeg",
    b"\x89PNG": "png",
    b"GIF8": "gif",
    b"BM": "bmp",
    b"II*\x00": "tiff",         # TIFF little-endian
    b"MM\x00*": "tiff",         # TIFF big-endian
    b"RIFF": "webp",             # WebP начинается с RIFF
}

ALLOWED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".txt", ".odt",
    ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"
}

def detect_magic(content: bytes) -> Optional[str]:
    """Определяет реальный тип файла по магическим байтам."""
    for magic, filetype in MAGIC_BYTES.items():
        if content[:len(magic)] == magic:
            return filetype
    # TXT не имеет магических байт — проверим, что это валидный текст
    try:
        content[:512].decode("utf-8")
        return "txt"
    except UnicodeDecodeError:
        try:
            content[:512].decode("cp1251")
            return "txt"
        except UnicodeDecodeError:
            return None

# ============================================================================
# ОСНОВНАЯ ФУНКЦИЯ ПАРСИНГА
# ============================================================================

async def parse_file(file: UploadFile) -> dict:
    """
    Главная функция парсинга файлов.
    Определяет тип файла и вызывает нужный парсер.
    
    Returns:
        dict с ключами:
        - text: извлечённый текст
        - file_type: тип файла
        - pages: количество страниц (если применимо)
        - success: True/False
        - error: описание ошибки (если есть)
    """
    filename = file.filename or ""
    ext = Path(filename.lower()).suffix

    # Проверка расширения
    if ext not in ALLOWED_EXTENSIONS:
        return {
            "text": "", "file_type": "unknown", "pages": 0, "success": False,
            "error": f"Формат файла не поддерживается: {filename}"
        }

    content = await file.read()

    # Проверка размера
    if len(content) > MAX_FILE_SIZE:
        return {
            "text": "", "file_type": "unknown", "pages": 0, "success": False,
            "error": f"Файл слишком большой: {len(content) // 1024 // 1024} МБ. Максимум — 10 МБ."
        }

    if len(content) == 0:
        return {
            "text": "", "file_type": "unknown", "pages": 0, "success": False,
            "error": "Файл пустой."
        }

    # Проверка реального типа по содержимому
    magic_type = detect_magic(content)

    # Сопоставляем расширение с ожидаемым магическим типом
    ext_to_magic = {
        ".pdf": {"pdf"},
        ".docx": {"zip"},
        ".doc": {"zip", None},   # старый .doc не имеет стандартных байт
        ".odt": {"zip"},
        ".txt": {"txt"},
        ".jpg": {"jpeg"}, ".jpeg": {"jpeg"},
        ".png": {"png"},
        ".tiff": {"tiff"}, ".tif": {"tiff"},
        ".bmp": {"bmp"},
        ".gif": {"gif"},
        ".webp": {"webp"},
    }

    allowed_magic = ext_to_magic.get(ext, set())
    if magic_type not in allowed_magic and ext != ".doc":
        return {
            "text": "", "file_type": "unknown", "pages": 0, "success": False,
            "error": f"Содержимое файла не соответствует расширению {ext}. Возможно, файл повреждён или переименован."
        }

    try:
        if ext == ".pdf":
            return await parse_pdf(content, filename)
        elif ext == ".docx":
            return await parse_docx(content, filename)
        elif ext == ".doc":
            return await parse_doc(content, filename)
        elif ext == ".txt":
            return await parse_txt(content, filename)
        elif ext == ".odt":
            return await parse_odt(content, filename)
        elif ext in (".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"):
            return await parse_image_ocr(content, filename)
        else:
            return {
                "text": "", "file_type": "unknown", "pages": 0, "success": False,
                "error": f"Формат файла не поддерживается: {filename}"
            }

    except Exception as e:
        return {
            "text": "", "file_type": "error", "pages": 0, "success": False,
            "error": f"Ошибка при обработке файла: {str(e)}"
        }

# ============================================================================
# PDF ПАРСЕР
# ============================================================================

async def parse_pdf(content: bytes, filename: str) -> dict:
    """Парсит PDF файл. Если текст не извлекается - применяет OCR."""
    
    try:
        import pymupdf  # PyMuPDF (fitz)
        
        doc = pymupdf.open(stream=content, filetype="pdf")
        text_parts = []
        pages = len(doc)
        
        for page_num in range(pages):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                text_parts.append(f"[Страница {page_num + 1}]\n{text}")
            else:
                # Страница скан — применяем OCR
                ocr_text = await ocr_page(page)
                if ocr_text:
                    text_parts.append(f"[Страница {page_num + 1} (OCR)]\n{ocr_text}")
        
        doc.close()
        
        full_text = "\n\n".join(text_parts)
        
        return {
            "text": full_text,
            "file_type": "pdf",
            "pages": pages,
            "success": True,
            "error": None
        }
    
    except ImportError:
        return await parse_pdf_fallback(content, filename)
    
    except Exception as e:
        return {
            "text": "", "file_type": "pdf", "pages": 0, "success": False,
            "error": f"Ошибка PDF: {str(e)}"
        }

async def parse_pdf_fallback(content: bytes, filename: str) -> dict:
    """Запасной PDF парсер через pypdf."""
    try:
        import pypdf
        
        reader = pypdf.PdfReader(io.BytesIO(content))
        text_parts = []
        pages = len(reader.pages)
        
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"[Страница {i + 1}]\n{text}")
        
        return {
            "text": "\n\n".join(text_parts),
            "file_type": "pdf",
            "pages": pages,
            "success": True,
            "error": None
        }
    
    except Exception as e:
        return {
            "text": "", "file_type": "pdf", "pages": 0, "success": False,
            "error": f"Ошибка PDF (fallback): {str(e)}"
        }

# ============================================================================
# WORD ПАРСЕР
# ============================================================================

async def parse_docx(content: bytes, filename: str) -> dict:
    """Парсит .docx файл (Word 2007+)."""
    
    try:
        from docx import Document
        
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Также извлекаем текст из таблиц
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        full_text = "\n".join(text_parts)
        
        return {
            "text": full_text,
            "file_type": "docx",
            "pages": len(doc.paragraphs),
            "success": True,
            "error": None
        }
    
    except Exception as e:
        return {
            "text": "", "file_type": "docx", "pages": 0, "success": False,
            "error": f"Ошибка DOCX: {str(e)}"
        }

async def parse_doc(content: bytes, filename: str) -> dict:
    """Парсит .doc файл (старый Word формат)."""
    
    tmp_path = None
    try:
        import textract
        
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        text = textract.process(tmp_path).decode('utf-8')
        
        return {
            "text": text,
            "file_type": "doc",
            "pages": 1,
            "success": True,
            "error": None
        }
    
    except Exception as e:
        return {
            "text": "", "file_type": "doc", "pages": 0, "success": False,
            "error": f"Ошибка DOC: {str(e)}. Попробуйте конвертировать в .docx"
        }
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

# ============================================================================
# ТЕКСТОВЫЙ ПАРСЕР
# ============================================================================

async def parse_txt(content: bytes, filename: str) -> dict:
    """Парсит .txt файл."""
    
    try:
        for encoding in ['utf-8', 'cp1251', 'latin-1']:
            try:
                text = content.decode(encoding)
                return {
                    "text": text,
                    "file_type": "txt",
                    "pages": 1,
                    "success": True,
                    "error": None
                }
            except UnicodeDecodeError:
                continue
        
        text = content.decode('utf-8', errors='replace')
        return {
            "text": text,
            "file_type": "txt",
            "pages": 1,
            "success": True,
            "error": None
        }
    
    except Exception as e:
        return {
            "text": "", "file_type": "txt", "pages": 0, "success": False,
            "error": f"Ошибка TXT: {str(e)}"
        }

# ============================================================================
# ODT ПАРСЕР
# ============================================================================

async def parse_odt(content: bytes, filename: str) -> dict:
    """Парсит .odt файл (OpenDocument)."""
    
    tmp_path = None
    try:
        from odf.opendocument import load
        from odf.text import P
        from odf import teletype
        
        with tempfile.NamedTemporaryFile(suffix='.odt', delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        doc = load(tmp_path)
        
        text_parts = []
        for para in doc.getElementsByType(P):
            text = teletype.extractText(para)
            if text.strip():
                text_parts.append(text)
        
        return {
            "text": "\n".join(text_parts),
            "file_type": "odt",
            "pages": 1,
            "success": True,
            "error": None
        }
    
    except Exception as e:
        return {
            "text": "", "file_type": "odt", "pages": 0, "success": False,
            "error": f"Ошибка ODT: {str(e)}"
        }
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

# ============================================================================
# OCR ПАРСЕР (для изображений и сканов)
# ============================================================================

async def parse_image_ocr(content: bytes, filename: str) -> dict:
    """Применяет OCR к изображению (JPG, PNG, TIFF и т.д.)."""
    
    try:
        import pytesseract
        from PIL import Image
        
        image = Image.open(io.BytesIO(content))
        
        # OCR на русском + казахском + английском языках
        text = pytesseract.image_to_string(
            image,
            lang='rus+kaz+eng',
            config='--psm 6'
        )
        
        return {
            "text": text,
            "file_type": "image_ocr",
            "pages": 1,
            "success": True,
            "error": None
        }
    
    except ImportError:
        return {
            "text": "", "file_type": "image", "pages": 0, "success": False,
            "error": "OCR не установлен. Выполните: pip install pytesseract pillow && установите Tesseract OCR"
        }
    
    except Exception as e:
        return {
            "text": "", "file_type": "image", "pages": 0, "success": False,
            "error": f"Ошибка OCR: {str(e)}"
        }

async def ocr_page(page) -> str:
    """Применяет OCR к странице PDF (для сканированных PDF)."""
    
    try:
        import pytesseract
        from PIL import Image
        
        pix = page.get_pixmap(dpi=300)
        img_data = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_data))
        
        text = pytesseract.image_to_string(
            image,
            lang='rus+kaz+eng',
            config='--psm 6'
        )
        
        return text
    
    except Exception:
        return ""

# ============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================================

def get_supported_formats() -> list:
    """Возвращает список поддерживаемых форматов."""
    return [
        ".pdf", ".docx", ".doc", ".txt", ".odt",
        ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".gif", ".webp"
    ]

def clean_text(text: str) -> str:
    """Очищает извлечённый текст от лишних символов."""
    
    if not text:
        return ""
    
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]
    cleaned = '\n'.join(lines)
    
    return cleaned

def truncate_text(text: str, max_chars: int = 10000) -> str:
    """Обрезает текст до максимальной длины для LLM."""
    
    if len(text) <= max_chars:
        return text
    
    return text[:max_chars] + f"\n\n[... текст обрезан, показаны первые {max_chars} символов из {len(text)}]"
